"""
Функции для генерации выходного файла с оформленным списком использованных источников.
"""
from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # pylint: disable=E0611
from docx.shared import Cm, Pt


class Renderer:
    """
    Создание выходного файла – Word.
    """

    def __init__(self, rows: tuple[str, ...]):
        self.rows = rows

    def render(self, path: Any) -> None:
        """
        Метод генерации Word-файла со списком использованных источников.

        :param Renderer self: Текущий объект
        :param Path path: Путь для сохранения выходного файла.
        :return:
        """

        document = Document()

        # стилизация заголовка
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runner = paragraph.add_run("Список использованной литературы")
        runner.bold = True

        # стилизация текста
        style_normal = document.styles["Normal"]
        style_normal.font.name = "Times New Roman"
        style_normal.font.size = Pt(12)
        style_normal.paragraph_format.line_spacing = 1.5
        style_normal.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        paragraph_format = self.get_styles(document)

        for row in self.rows:
            # добавление источника
            if "ITALIC" in row:
                parts = row.split("ITALIC")
                part = document.add_paragraph(parts[0], style=paragraph_format)
                part.add_run(parts[1]).italic = True
                part.add_run(parts[2])

            else:
                document.add_paragraph(row, style=paragraph_format)

        # сохранение файла Word
        document.save(path)

    def get_styles(self, document: Any) -> Any:
        """
        Метод получения стилей.

        :param document | Any document: Текущий документ.
        :return:
        """
        style_normal = document.styles["Normal"]
        return style_normal


class GOSTRenderer(Renderer):
    def get_styles(self, document: Any) -> Any:
        """
        Метод получения стилей. Для ГОСТ - Нумерованный список.
        :param document | Any document: Текущий документ.
        :return: Наименование стиля абзаца для ГОСТ.
        """
        style_normal = document.styles["List Number"]
        return style_normal


class APARenderer(Renderer):
    def get_styles(self, document: Any) -> None:
        """
        Метод получения стилей. Для APA - НЕ нумерованный список и отступ.
        :param document | Any document: Текущий документ.
        """
        style_normal = document.styles["Normal"]
        style_normal.paragraph_format.first_line_indent = Cm(-1.5)
        return style_normal
